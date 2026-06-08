import streamlit as st
import pandas as pd
import json
import re
import io
import os
from pathlib import Path
import sqlite3
from datetime import datetime
from dotenv import load_dotenv

load_dotenv()

# ── Optional imports with graceful fallback ───────────────
try:
    import pdfplumber
    PDF_OK = True
except ImportError:
    PDF_OK = False

try:
    import docx
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    from PIL import Image
    import base64
    PIL_OK = True
except ImportError:
    PIL_OK = False

try:
    import pytesseract
    OCR_OK = True
except ImportError:
    OCR_OK = False

try:
    from openai import OpenAI
    OPENAI_OK = True
except ImportError:
    OPENAI_OK = False

st.set_page_config(
    page_title="Upload - Finteca AuditRep",
    page_icon="🏦",
    layout="wide"
)

CSS = """
<style>
.finteca-header {
    background: linear-gradient(135deg, #1a237e 0%, #283593 50%, #42a5f5 100%);
    padding: 25px 30px; border-radius: 12px; color: white;
    margin-bottom: 25px; box-shadow: 0 4px 15px rgba(26,35,126,0.3);
}
.finteca-header h1 { margin:0; font-size:2.2em; font-weight:800; }
.finteca-header p  { margin:5px 0 0 0; opacity:0.85; }
.finteca-badge {
    background:rgba(255,255,255,0.2); padding:3px 10px;
    border-radius:20px; font-size:0.75em; display:inline-block; margin-top:8px;
}
.section-header {
    background:#f5f7ff; border-left:4px solid #1a237e; padding:10px 15px;
    border-radius:0 8px 8px 0; margin:15px 0; font-weight:600; color:#1a237e;
}
.alert-high {
    background:#fff3e0; border-left:5px solid #e65100;
    padding:12px 15px; border-radius:6px; margin:6px 0;
}
.alert-ok {
    background:#e8f5e9; border-left:5px solid #2e7d32;
    padding:12px 15px; border-radius:6px; margin:6px 0;
}
.finteca-footer {
    text-align:center; color:#999; font-size:0.8em;
    padding:20px; border-top:1px solid #eee; margin-top:30px;
}
</style>
"""
st.markdown(CSS, unsafe_allow_html=True)

st.markdown("""
<div class="finteca-header">
    <h1>🏦 Finteca AuditRep</h1>
    <p>Document Upload & Processing Centre</p>
    <span class="finteca-badge">Module 1 — Upload</span>
</div>
""", unsafe_allow_html=True)

# ── Show module status ────────────────────────────────────
with st.expander("🔧 System Status", expanded=False):
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("PDF Support",   "✅ Ready" if PDF_OK   else "❌ Install pdfplumber")
    c2.metric("Word Support",  "✅ Ready" if DOCX_OK  else "❌ Install python-docx")
    c3.metric("Image Support", "✅ Ready" if PIL_OK   else "❌ Install pillow")
    c4.metric("OCR Support",   "✅ Ready" if OCR_OK   else "❌ Install pytesseract")
    c5.metric("AI Support",    "✅ Ready" if OPENAI_OK else "❌ Install openai")

# Cloud + Local compatible
DB_PATH = "/tmp/reconciliation.db" if os.path.exists("/mount/src") else "data/reconciliation.db"
if not os.path.exists("/mount/src"):
    Path("data").mkdir(exist_ok=True)
API_KEY    = os.getenv("OPENAI_API_KEY", "")
TABLE_LIST = [
    "purchases", "sales", "banking", "collections",
    "sales_returns", "swap_deals", "inventory"
]
TYPE_TO_TABLE = {
    "purchases": "purchases", "purchase_order": "purchases",
    "sales": "sales", "sales_invoice": "sales", "invoice": "sales",
    "banking": "banking", "bank_statement": "banking",
    "collections": "collections", "collection": "collections",
    "receipt": "collections",
    "sales_returns": "sales_returns", "sales_return": "sales_returns",
    "return": "sales_returns",
    "swap_deals": "swap_deals", "swap": "swap_deals",
    "inventory": "inventory",
}

# ── Database helpers ──────────────────────────────────────
def get_conn():
    return sqlite3.connect(DB_PATH, check_same_thread=False)

def clean_df(df):
    if df is None or df.empty:
        return pd.DataFrame()
    df = df.dropna(how="all")
    df = df.loc[:, df.notna().any()]
    df.columns = [
        re.sub(r"[^a-zA-Z0-9_]", "_", str(c).strip().lower())
        for c in df.columns
    ]
    for col in df.columns:
        try:
            cleaned = (
                df[col].astype(str)
                .str.replace(",", "", regex=False)
                .str.replace("$", "", regex=False)
                .str.strip()
            )
            num = pd.to_numeric(cleaned, errors="coerce")
            if num.notna().sum() > len(df) * 0.4:
                df[col] = num
        except Exception:
            pass
    return df.reset_index(drop=True)

def get_table_cols(table):
    conn = get_conn()
    cursor = conn.cursor()
    cursor.execute(f"PRAGMA table_info({table})")
    cols = [r[1] for r in cursor.fetchall()]
    conn.close()
    return cols

def map_and_save(df, table, filename):
    try:
        conn       = get_conn()
        valid_cols = get_table_cols(table)
        col_map = {
            "date": ["date","trans_date","transaction_date","invoice_date",
                     "doc_date","tran_date","posting_date","value_date"],
            "reference_number": ["reference","ref","ref_no","invoice_no",
                                  "po_number","voucher_no","receipt_no",
                                  "bill_no","document_no","order_no"],
            "invoice_number": ["invoice_no","invoice_number","inv_no",
                                "invoice","bill_no","receipt_no"],
            "supplier": ["supplier","vendor","creditor","from",
                          "supplier_name","vendor_name"],
            "customer": ["customer","client","debtor","buyer",
                          "customer_name","client_name","sold_to"],
            "item_description": ["description","item","product","goods",
                                  "particulars","details","narration",
                                  "item_name","item_description"],
            "quantity": ["qty","quantity","units","pieces",
                          "no_of_items","count","nos","num"],
            "unit_cost": ["unit_cost","cost_price","purchase_price","cost_per_unit"],
            "unit_price": ["unit_price","selling_price","price","sale_price"],
            "total_cost": ["total_cost","total","amount","value",
                            "total_amount","gross_amount","line_total"],
            "gross_amount": ["gross","gross_amount","subtotal","sub_total"],
            "discount": ["discount","disc","rebate","reduction"],
            "tax": ["tax","vat","gst","sales_tax","tax_amount"],
            "net_amount": ["net","net_amount","net_total","invoice_total",
                            "total_due","amount_due","payable"],
            "debit": ["debit","dr","withdrawal","amount_out","out","payments"],
            "credit": ["credit","cr","deposit","receipt","amount_in","in","receipts"],
            "balance": ["balance","running_balance","closing_balance","ledger_balance"],
            "payment_method": ["payment_method","mode","payment_mode",
                                "method","pay_mode","payment_type"],
            "payment_reference": ["payment_ref","payment_reference","bank_ref",
                                   "cheque_no","transfer_ref","chq_no"],
            "amount": ["amount","value","total","sum","net_amount","receipt_amount"],
            "description": ["description","narration","details",
                             "particulars","remarks","transaction_details"],
        }

        renamed = df.copy()
        renamed.columns = [
            re.sub(r"[^a-zA-Z0-9_]", "_", str(c).strip().lower())
            for c in renamed.columns
        ]

        rename_dict  = {}
        used_targets = set()
        for target, patterns in col_map.items():
            if target in valid_cols and target not in used_targets:
                for col in renamed.columns:
                    if col.strip("_") in patterns:
                        rename_dict[col] = target
                        used_targets.add(target)
                        break

        renamed = renamed.rename(columns=rename_dict)
        renamed["document_source"] = filename
        renamed["uploaded_at"]     = datetime.now().isoformat()

        save_cols = [c for c in renamed.columns if c in valid_cols]
        save_df   = renamed[save_cols] if save_cols else renamed

        before = pd.read_sql(
            f"SELECT COUNT(*) as c FROM {table}", conn
        )["c"].iloc[0]

        save_df.to_sql(table, conn, if_exists="append", index=False)

        after = pd.read_sql(
            f"SELECT COUNT(*) as c FROM {table}", conn
        )["c"].iloc[0]

        saved = int(after - before)

        conn.execute(
            "INSERT INTO upload_log "
            "(filename, document_type, rows_extracted, rows_saved, status) "
            "VALUES (?,?,?,?,?)",
            (filename, table, len(df), saved, "success"),
        )
        conn.commit()
        conn.close()
        return {"success": True, "rows_saved": saved}

    except Exception as e:
        return {"success": False, "error": str(e)}

def ai_classify(text, filename):
    if not API_KEY or not OPENAI_OK:
        return {
            "document_type": "unknown",
            "confidence": 0,
            "red_flags": [],
            "notes": "No API key or OpenAI not installed",
        }
    try:
        client = OpenAI(api_key=API_KEY)
        resp   = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {
                    "role": "system",
                    "content": (
                        "You are Finteca AuditRep AI. "
                        "Classify financial documents. Return JSON only."
                    ),
                },
                {
                    "role": "user",
                    "content": (
                        f"Classify this financial document.\n"
                        f"Filename: {filename}\n"
                        f"Content sample: {text[:2000]}\n\n"
                        "Return JSON:\n"
                        "{\n"
                        '  "document_type": "purchases|sales|banking|collections|sales_returns|swap_deals|inventory|other",\n'
                        '  "confidence": 0-100,\n'
                        '  "red_flags": [],\n'
                        '  "notes": ""\n'
                        "}"
                    ),
                },
            ],
            max_tokens=400,
            response_format={"type": "json_object"},
        )
        return json.loads(resp.choices[0].message.content)
    except Exception as e:
        return {
            "document_type": "unknown",
            "confidence": 0,
            "red_flags": [],
            "notes": str(e),
        }

def process_file(uploaded_file):
    name     = uploaded_file.name
    ext      = name.rsplit(".", 1)[-1].lower()
    sheets   = {}
    raw_text = ""

    try:
        # ── Excel / CSV ───────────────────────────────────
        if ext in ["xlsx", "xls"]:
            xl = pd.ExcelFile(uploaded_file)
            for sh in xl.sheet_names:
                df = clean_df(pd.read_excel(uploaded_file, sheet_name=sh))
                if not df.empty:
                    sheets[sh] = df
            raw_text = " ".join(
                str(c) for df in sheets.values() for c in df.columns
            )

        elif ext == "csv":
            df = clean_df(pd.read_csv(uploaded_file))
            if not df.empty:
                sheets["Data"] = df
            raw_text = " ".join(df.columns.tolist()) if not df.empty else ""

        # ── PDF ───────────────────────────────────────────
        elif ext == "pdf":
            if not PDF_OK:
                st.error("pdfplumber not installed. Run: pip install pdfplumber")
                return {"sheets": {}, "classification": {}, "total_rows": 0}

            with pdfplumber.open(uploaded_file) as pdf:
                for i, page in enumerate(pdf.pages):
                    txt       = page.extract_text() or ""
                    raw_text += txt
                    for tbl in page.extract_tables():
                        if tbl and len(tbl) > 1:
                            try:
                                df = clean_df(
                                    pd.DataFrame(tbl[1:], columns=tbl[0])
                                )
                                if not df.empty:
                                    sheets[f"Page{i+1}_Table"] = df
                            except Exception:
                                pass

        # ── Word ──────────────────────────────────────────
        elif ext in ["docx", "doc"]:
            if not DOCX_OK:
                st.error("python-docx not installed. Run: pip install python-docx")
                return {"sheets": {}, "classification": {}, "total_rows": 0}

            d        = docx.Document(uploaded_file)
            raw_text = "\n".join(p.text for p in d.paragraphs)
            for i, tbl in enumerate(d.tables):
                data = [
                    [c.text.strip() for c in r.cells]
                    for r in tbl.rows
                ]
                if data and len(data) > 1:
                    df = clean_df(
                        pd.DataFrame(data[1:], columns=data[0])
                    )
                    if not df.empty:
                        sheets[f"Table{i+1}"] = df

        # ── Images ────────────────────────────────────────
        elif ext in ["png", "jpg", "jpeg", "tiff", "bmp"]:
            if not PIL_OK:
                st.error("pillow not installed. Run: pip install pillow")
                return {"sheets": {}, "classification": {}, "total_rows": 0}

            img = Image.open(uploaded_file)

            if OCR_OK:
                raw_text = pytesseract.image_to_string(img)

            if API_KEY and OPENAI_OK and PIL_OK:
                try:
                    buf = io.BytesIO()
                    img.save(buf, format="PNG")
                    b64 = base64.b64encode(buf.getvalue()).decode()

                    client = OpenAI(api_key=API_KEY)
                    resp   = client.chat.completions.create(
                        model="gpt-4o",
                        messages=[{
                            "role": "user",
                            "content": [
                                {
                                    "type": "image_url",
                                    "image_url": {
                                        "url": f"data:image/png;base64,{b64}"
                                    },
                                },
                                {
                                    "type": "text",
                                    "text": (
                                        "Extract all financial data as JSON:\n"
                                        '{"document_type":"","date":"","reference":"","party":"",'
                                        '"line_items":[{"description":"","quantity":0,"unit_price":0,"total":0}],'
                                        '"totals":{"net_total":0},"red_flags":[]}'
                                    ),
                                },
                            ],
                        }],
                        max_tokens=1000,
                    )
                    content = resp.choices[0].message.content
                    m = re.search(r"\{.*\}", content, re.DOTALL)
                    if m:
                        data  = json.loads(m.group())
                        items = data.get("line_items", [])
                        if items:
                            sheets["Extracted"] = pd.DataFrame(items)
                        raw_text += str(data)
                except Exception:
                    pass

    except Exception as e:
        return {
            "sheets": {},
            "classification": {
                "document_type": "error",
                "confidence": 0,
                "red_flags": [str(e)],
                "notes": str(e),
            },
            "total_rows": 0,
        }

    classification = ai_classify(raw_text, name)
    return {
        "sheets":         sheets,
        "classification": classification,
        "total_rows":     sum(len(d) for d in sheets.values()),
    }

# ── UI Layout ─────────────────────────────────────────────
col_up, col_man = st.columns([3, 2])

with col_up:
    st.markdown(
        '<div class="section-header">📁 Upload Documents</div>',
        unsafe_allow_html=True,
    )
    uploaded_files = st.file_uploader(
        "Drop files here — Excel, CSV, PDF, Word, Images",
        accept_multiple_files=True,
        type=[
            "xlsx", "xls", "csv",
            "pdf", "docx", "doc",
            "png", "jpg", "jpeg", "tiff", "bmp",
        ],
    )

with col_man:
    st.markdown(
        '<div class="section-header">✏️ Manual Entry</div>',
        unsafe_allow_html=True,
    )
    entry_type = st.selectbox(
        "Entry Type",
        ["Sale", "Purchase", "Collection",
         "Bank Credit", "Bank Debit", "Sales Return"],
    )
    with st.form("manual_form", clear_on_submit=True):
        m_date    = st.date_input("Date")
        m_ref     = st.text_input("Reference / Invoice No.")
        m_party   = st.text_input("Customer / Supplier")
        m_desc    = st.text_input("Description / Item")
        m_qty     = st.number_input("Quantity",  min_value=0.0, value=1.0)
        m_amount  = st.number_input("Amount",    min_value=0.0, value=0.0)
        m_payment = st.selectbox(
            "Payment Method",
            ["Cash", "Bank Transfer", "Cheque", "Card", "Credit", "Other"],
        )

        if st.form_submit_button(
            "➕ Add Entry", type="primary", use_container_width=True
        ):
            tmap = {
                "Sale": (
                    "sales",
                    {
                        "date": str(m_date), "invoice_number": m_ref,
                        "customer": m_party, "item_description": m_desc,
                        "quantity": m_qty, "net_amount": m_amount,
                        "payment_method": m_payment,
                    },
                ),
                "Purchase": (
                    "purchases",
                    {
                        "date": str(m_date), "reference_number": m_ref,
                        "supplier": m_party, "item_description": m_desc,
                        "quantity": m_qty, "total_cost": m_amount,
                        "payment_method": m_payment,
                    },
                ),
                "Collection": (
                    "collections",
                    {
                        "date": str(m_date), "invoice_reference": m_ref,
                        "customer": m_party, "amount": m_amount,
                        "payment_method": m_payment,
                    },
                ),
                "Bank Credit": (
                    "banking",
                    {
                        "date": str(m_date), "reference": m_ref,
                        "description": m_desc, "credit": m_amount,
                        "debit": 0, "transaction_type": "credit",
                    },
                ),
                "Bank Debit": (
                    "banking",
                    {
                        "date": str(m_date), "reference": m_ref,
                        "description": m_desc, "debit": m_amount,
                        "credit": 0, "transaction_type": "debit",
                    },
                ),
                "Sales Return": (
                    "sales_returns",
                    {
                        "date": str(m_date), "return_reference": m_ref,
                        "customer": m_party, "item_description": m_desc,
                        "quantity_returned": m_qty, "return_amount": m_amount,
                    },
                ),
            }

            if entry_type in tmap:
                tbl, row_data = tmap[entry_type]
                res = map_and_save(pd.DataFrame([row_data]), tbl, "manual_entry")
                if res["success"]:
                    st.success(f"✅ Saved to {tbl}!")
                    st.rerun()
                else:
                    st.error(f"❌ {res.get('error')}")

# ── Process uploaded files ────────────────────────────────
if uploaded_files:
    st.divider()
    st.markdown(
        '<div class="section-header">🔍 Extraction Results</div>',
        unsafe_allow_html=True,
    )

    for uf in uploaded_files:
        with st.spinner(f"🤖 Processing: {uf.name}"):
            result = process_file(uf)

        clf   = result.get("classification", {})
        dtype = clf.get("document_type", "unknown")
        conf  = clf.get("confidence", 0)
        flags = clf.get("red_flags", [])
        total = result.get("total_rows", 0)
        icon  = "🟢" if conf >= 80 else "🟡" if conf >= 50 else "🔴"

        with st.expander(
            f"{icon} {uf.name}  |  {dtype.upper()}  "
            f"|  {conf}% confidence  |  {total} rows",
            expanded=True,
        ):
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Document Type",  dtype.replace("_", " ").title())
            c2.metric("AI Confidence",  f"{conf}%")
            c3.metric("Rows Extracted", total)
            c4.metric("Red Flags",      len(flags))

            for flag in flags:
                st.markdown(
                    f'<div class="alert-high">🚩 {flag}</div>',
                    unsafe_allow_html=True,
                )

            if clf.get("notes"):
                st.info(f"💡 {clf['notes']}")

            sheets = result.get("sheets", {})
            for sh_name, df in sheets.items():
                if not df.empty:
                    st.markdown(
                        f"**📋 {sh_name}** — "
                        f"{len(df)} rows × {len(df.columns)} columns"
                    )
                    st.dataframe(df.head(10), use_container_width=True, height=200)

            if sheets:
                suggested = TYPE_TO_TABLE.get(dtype, "sales")
                sc1, sc2  = st.columns([2, 1])
                with sc1:
                    sel_table = st.selectbox(
                        "Save to table:",
                        TABLE_LIST,
                        index=(
                            TABLE_LIST.index(suggested)
                            if suggested in TABLE_LIST
                            else 0
                        ),
                        key=f"tbl_{uf.name}",
                    )
                with sc2:
                    sel_sheet = st.selectbox(
                        "Sheet:",
                        list(sheets.keys()),
                        key=f"sh_{uf.name}",
                    )

                if st.button(
                    "💾 Save to database",
                    key=f"save_{uf.name}",
                    type="primary",
                    use_container_width=True,
                ):
                    save_df = sheets.get(sel_sheet, pd.DataFrame())
                    if not save_df.empty:
                        res = map_and_save(save_df, sel_table, uf.name)
                        if res["success"]:
                            st.success(
                                f"✅ {res['rows_saved']} rows saved "
                                f"to **{sel_table}**!"
                            )
                        else:
                            st.error(f"❌ {res.get('error')}")
                    else:
                        st.warning("No data to save.")

# ── Database Status ───────────────────────────────────────
st.divider()
st.markdown(
    '<div class="section-header">📊 Database Status</div>',
    unsafe_allow_html=True,
)
icons  = ["🛒", "💰", "🏦", "💵", "↩️", "🔄", "📦"]
tables = [
    "purchases", "sales", "banking", "collections",
    "sales_returns", "swap_deals", "inventory",
]
conn2 = get_conn()
cols  = st.columns(7)
for col, icon, tbl in zip(cols, icons, tables):
    try:
        n = int(
            pd.read_sql(f"SELECT COUNT(*) as c FROM {tbl}", conn2)["c"].iloc[0]
        )
    except Exception:
        n = 0
    col.metric(f"{icon} {tbl.replace('_', ' ').title()}", n)
conn2.close()

try:
    log_conn = get_conn()
    log = pd.read_sql(
        "SELECT filename, document_type, rows_saved, status, uploaded_at "
        "FROM upload_log ORDER BY uploaded_at DESC LIMIT 20",
        log_conn,
    )
    log_conn.close()
    if not log.empty:
        st.markdown(
            '<div class="section-header">📜 Upload History</div>',
            unsafe_allow_html=True,
        )
        st.dataframe(log, use_container_width=True, height=200)
except Exception:
    pass

st.markdown(
    '<div class="finteca-footer">Finteca AuditRep v1.0.0</div>',
    unsafe_allow_html=True,
)
